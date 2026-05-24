from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph


SOURCE = Path("outputs/Диплом_AdoptPet_актуалізована_з_новими_діаграмами.docx")
OUTPUT = Path("outputs/Диплом_AdoptPet_фінальна_з_діаграмами_та_скріншотами.docx")
SCREEN_DIR = Path("outputs/diploma_ui_screenshots")
SCREENS = {
    "home": SCREEN_DIR / "01_home.png",
    "catalog": SCREEN_DIR / "02_catalog.png",
    "lost_found": SCREEN_DIR / "03_lost_found.png",
    "donation": SCREEN_DIR / "04_donation.png",
    "profile": SCREEN_DIR / "05_profile.png",
    "dashboard": SCREEN_DIR / "06_dashboard.png",
    "login": SCREEN_DIR / "07_login.png",
    "assistant": SCREEN_DIR / "08_ai_assistant.png",
    "lost_found_form": SCREEN_DIR / "09_lost_found_form.png",
    "animal_details": SCREEN_DIR / "10_animal_details.png",
    "adoption_form": SCREEN_DIR / "11_adoption_form.png",
    "admin_top": SCREEN_DIR / "12_admin_panel_top.png",
    "admin_lists": SCREEN_DIR / "13_admin_panel_lists.png",
}


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def find_index(doc: Document, prefix: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index
    raise ValueError(f"Paragraph not found: {prefix}")


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def clone_empty(paragraph: Paragraph, before: bool) -> Paragraph:
    element = deepcopy(paragraph._p)
    element.clear_content()
    if before:
        paragraph._p.addprevious(element)
    else:
        paragraph._p.addnext(element)
    result = Paragraph(element, paragraph._parent)
    if paragraph.style:
        result.style = paragraph.style
    return result


def insert_before(paragraph: Paragraph, text: str = "") -> Paragraph:
    result = clone_empty(paragraph, before=True)
    if text:
        result.add_run(text)
    return result


def add_picture(paragraph: Paragraph, image_path: Path, width_inches: float) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def replace_image_before_caption(
    doc: Document, caption_prefix: str, image_path: Path, width_inches: float = 6.35
) -> Paragraph:
    caption_index = find_index(doc, caption_prefix)
    image_paragraph = doc.paragraphs[caption_index - 1]
    add_picture(image_paragraph, image_path, width_inches)
    return doc.paragraphs[caption_index]


def insert_figure_before(
    anchor: Paragraph,
    image_path: Path,
    caption: str,
    explanation: str,
    width_inches: float = 6.35,
) -> None:
    image_paragraph = insert_before(anchor)
    add_picture(image_paragraph, image_path, width_inches)
    caption_paragraph = insert_before(anchor, caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_before(anchor, explanation)


def replace_existing_figures(doc: Document) -> None:
    caption = replace_image_before_caption(doc, "Рис. 5.1.", SCREENS["home"])
    set_text(caption, "Рис. 5.1. Головна сторінка платформи AdoptPet")

    caption = replace_image_before_caption(doc, "Рис. 5.2.", SCREENS["catalog"])
    set_text(caption, "Рис. 5.2. Каталог адопції з пошуком і фільтрацією")

    caption = replace_image_before_caption(
        doc, "Рис. 5.3.", SCREENS["animal_details"]
    )
    set_text(caption, "Рис. 5.7. Детальна картка тварини для адопції")

    caption = find_paragraph(doc, "Рис. 5.4.")
    set_text(caption, "Рис. 5.9. Тестування REST API ML-модуля через Swagger UI")

    caption = replace_image_before_caption(doc, "Рис. 5.5.", SCREENS["admin_top"])
    set_text(caption, "Рис. 5.11. Адміністративна панель керування платформою")

    caption = replace_image_before_caption(doc, "Рис. 5.6.", SCREENS["assistant"], 2.7)
    set_text(caption, "Рис. 5.14. Діалогове вікно AI-помічника платформи AdoptPet")

    set_text(
        find_paragraph(doc, "Рис. 5.7. Результати"),
        "Рис. 5.15. Результати виконання автоматизованих тестів",
    )
    set_text(
        find_paragraph(doc, "Рис. 5.8. Динаміка"),
        "Рис. 5.16. Динаміка train/val Top-1 accuracy на датасетах Kaggle, Oxford та Stanford",
    )
    set_text(
        find_paragraph(doc, "Рис. 5.9. Діаграма"),
        "Рис. 5.17. Діаграма розгортання системи AdoptPet",
    )


def add_user_flow_figures(doc: Document) -> None:
    anchor = find_paragraph(doc, "У процесі тестування перевірялося:")
    insert_before(
        anchor,
        "Головна сторінка забезпечує швидкий перехід до двох ключових "
        "сценаріїв: адопції та пошуку загублених або знайдених тварин. "
        "У каталозі перевірено відображення карток, статусу доступності, "
        "пошуку та комбінації фільтрів за типом, віком і розміром.",
    )
    insert_figure_before(
        anchor,
        SCREENS["lost_found"],
        "Рис. 5.3. Розділ оголошень про загублених і знайдених тварин",
        "Модуль lost/found містить окремий інформаційний екран, перемикач "
        "типу оголошень і пошук за даними тварини та містом. На сторінці "
        "виділено дію створення нового оголошення користувачем.",
    )
    insert_figure_before(
        anchor,
        SCREENS["lost_found_form"],
        "Рис. 5.4. Публічна форма створення lost/found оголошення",
        "Форма дозволяє вказати тип оголошення, характеристики тварини, "
        "місце зникнення або знахідки й додати фото з телефону чи комп'ютера. "
        "Надіслане користувачем оголошення надходить на адміністративну модерацію.",
    )
    insert_figure_before(
        anchor,
        SCREENS["donation"],
        "Рис. 5.5. Сторінка разової допомоги тваринам",
        "Окрема сторінка допомоги розширює призначення платформи: користувач "
        "може обрати суму внеску та напрям підтримки. Інтерфейс підготовлений "
        "до подальшого підключення платіжного провайдера або реквізитів фонду.",
    )
    insert_figure_before(
        anchor,
        SCREENS["profile"],
        "Рис. 5.6. Особистий кабінет і завантаження фото профілю",
        "В особистому кабінеті відображаються дані облікового запису та "
        "зведені показники користувача. Повторно використовуваний компонент "
        "завантаження підтримує вибір локального фото або введення URL, "
        "а зображення зберігається через Cloudinary.",
    )

    animal_anchor = find_paragraph(doc, "У межах тестування перевірялися:")
    insert_before(
        animal_anchor,
        "Детальна сторінка тварини показує основну інформацію, статус, "
        "AI-аналіз породи та згортаний блок медичного стану. Це дозволяє "
        "не перевантажувати публічну картку, але зберегти доступ до важливих "
        "відомостей перед поданням заявки.",
    )
    insert_figure_before(
        animal_anchor,
        SCREENS["adoption_form"],
        "Рис. 5.8. Форма подання заявки на адопцію",
        "З детальної картки користувач переходить до форми заявки, де вводить "
        "контактні дані, адресу, досвід утримання тварин і додаткову "
        "інформацію. Подана форма зберігається для подальшого розгляду адміністратором.",
    )


def add_access_and_admin_figures(doc: Document) -> None:
    auth_anchor = find_paragraph(doc, "У процесі тестування перевірялися такі сценарії:")
    insert_figure_before(
        auth_anchor,
        SCREENS["login"],
        "Рис. 5.10. Сторінка входу та реєстрації користувача",
        "Для доступу до персональних сценаріїв користувач може увійти за "
        "обліковими даними або через Google-провайдер. Перевірено, що "
        "навігація змінюється відповідно до ролі та стану сесії.",
    )

    admin_anchor = find_paragraph(doc, "У межах тестування адміністративної частини")
    insert_before(
        admin_anchor,
        "Адміністративна панель об'єднує керування оголошеннями, створення "
        "нових записів і перехід до аналітичного dashboard. Для списків "
        "передбачені пошук, зміна статусу, модерація та згортання блоків, "
        "що залишається зручним при збільшенні кількості записів.",
    )
    insert_figure_before(
        admin_anchor,
        SCREENS["admin_lists"],
        "Рис. 5.12. Опрацювання заявок і згортані картки Animal в адмінпанелі",
        "У нижній частині адмінпанелі адміністратор переглядає заявки на "
        "адопцію, схвалює або відхиляє їх і відкриває внутрішні картки Animal. "
        "Картки містять службові дані про стан тварини, документи та обставини знахідки.",
    )
    insert_figure_before(
        admin_anchor,
        SCREENS["dashboard"],
        "Рис. 5.13. Dashboard зі зведеною аналітикою тварин",
        "Dashboard відображає кількість тварин у базі, активні анкети, "
        "кількість адопцій, lost/found оголошень, користувачів і заявок. "
        "Діаграми за видами та статусами створюють основу для управлінського "
        "аналізу, а нижчі блоки інтерфейсу доповнюють його породами, віком і регіонами.",
    )

    ai_anchor = find_paragraph(doc, "Також перевірялася робота негативних сценаріїв")
    insert_before(
        ai_anchor,
        "AI-помічник представлений у вигляді доступного з будь-якої сторінки "
        "чату. Користувач формулює потребу природною мовою, а система "
        "підбирає релевантні анкети з каталогу та повертає рекомендацію українською мовою.",
    )


def update_section_text(doc: Document) -> None:
    set_text(
        find_paragraph(doc, "На першому етапі було виконано перевірку клієнтської"),
        "На першому етапі було виконано перевірку користувацького інтерфейсу "
        "та наскрізних сценаріїв: переходу з головної сторінки до каталогу, "
        "перегляду картки тварини, подання заявки, створення lost/found "
        "оголошення, завантаження фото й роботи особистого кабінету.",
    )
    set_text(
        find_paragraph(doc, "Наступним етапом було тестування сторінки профілю тварини"),
        "Наступним етапом було тестування детальної картки тварини та "
        "пов'язаного процесу адопції. Сторінка містить фото, характеристики, "
        "опис, статус, блок медичної інформації, AI-аналіз породи й кнопку "
        "переходу до заявки.",
    )
    set_text(
        find_paragraph(doc, "Окрему увагу приділено тестуванню системи авторизації"),
        "Окрему увагу приділено тестуванню реєстрації, авторизації та "
        "розмежування прав доступу. Після входу користувач отримує особистий "
        "кабінет, а користувач із роллю адміністратора також отримує доступ "
        "до панелі керування і dashboard-аналітики.",
    )
    set_text(
        find_paragraph(doc, "У межах тестування адміністративної частини перевірялися:"),
        "У межах тестування адміністративної частини перевірялися:",
    )


def main() -> None:
    missing = [path for path in [SOURCE, *SCREENS.values()] if not path.exists()]
    if missing:
        raise FileNotFoundError("\n".join(str(path) for path in missing))

    doc = Document(SOURCE)
    replace_existing_figures(doc)
    update_section_text(doc)
    add_user_flow_figures(doc)
    add_access_and_admin_figures(doc)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())
    print(f"paragraphs={len(doc.paragraphs)}, inline_shapes={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
