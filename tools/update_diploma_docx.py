from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\Huawei\Downloads\Диплом.docx")
OUTPUT = Path("outputs/Диплом_оновлена_версія_AdoptPet.docx")


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_after(paragraph: Paragraph, text: str, style: str | None = None, bold_prefix: str | None = None) -> Paragraph:
    new_element = deepcopy(paragraph._p)
    new_element.clear_content()
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if style:
        new_paragraph.style = style
    elif paragraph.style:
        new_paragraph.style = paragraph.style
    if bold_prefix and text.startswith(bold_prefix):
        run = new_paragraph.add_run(bold_prefix)
        run.bold = True
        new_paragraph.add_run(text[len(bold_prefix):])
    else:
        new_paragraph.add_run(text)
    return new_paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def remove_exact_paragraphs(doc: Document, texts: set[str]) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in texts:
            remove_paragraph(paragraph)


def bullet_after(paragraph: Paragraph, items: list[str]) -> Paragraph:
    current = paragraph
    for item in items:
        current = insert_after(current, f"•\t{item}", style=paragraph.style.name)
    return current


def update_annotation(doc: Document) -> None:
    p = find_paragraph(doc, "У роботі спроєктовано та реалізовано fullstack-застосунок")
    set_text(
        p,
        "У роботі спроєктовано та реалізовано fullstack-застосунок на основі Next.js, "
        "MongoDB Atlas та Auth.js. Система підтримує адопцію тварин, каталог із пошуком, "
        "фільтрацією та пагінацією, подання заявок, модуль загублених і знайдених тварин, "
        "особистий кабінет користувача, адміністративну панель та AI-помічника для "
        "персоналізованого підбору тварин.",
    )
    insert_after(
        p,
        "Додатково реалізовано окрему модель Animal для внутрішнього обліку тварин, "
        "збереження медичного стану, документів, інформації про місце та особу, яка знайшла "
        "тварину. Для роботи із зображеннями інтегровано Cloudinary, що дозволяє користувачам "
        "і адміністраторам завантажувати фото з комп’ютера або мобільного пристрою. "
        "Адміністративна частина доповнена dashboard-аналітикою, пошуком, фільтрацією та "
        "модерацією користувацьких оголошень.",
        style=p.style.name,
    )
    set_text(
        find_paragraph(doc, "Проведено тестування функціональності системи"),
        "Проведено тестування функціональності системи, інтеграції ML-модуля, завантаження "
        "зображень, роботи особистого кабінету, адміністративних сценаріїв, модерації "
        "оголошень та процесу контейнеризації із використанням Docker.",
    )
    set_text(
        find_paragraph(doc, "Практичне значення роботи полягає у створенні програмного комплексу"),
        "Практичне значення роботи полягає у створенні програмного комплексу, який може бути "
        "використаний як основа для сучасної цифрової платформи адопції, обліку тварин, "
        "обробки заявок, пошуку загублених і знайдених тварин та підтримки управлінських "
        "рішень за допомогою аналітики.",
    )
    set_text(
        find_paragraph(doc, "Ключові слова: AdoptPet"),
        "Ключові слова: AdoptPet, веб-застосунок, адопція тварин, загублені та знайдені "
        "тварини, адміністративна панель, аналітика, Cloudinary, Auth.js, Next.js, MongoDB, "
        "Mongoose, штучний інтелект, машинне навчання, комп’ютерний зір, EfficientNet, FastAPI.",
    )
    set_text(
        find_paragraph(doc, "Сторінок 106"),
        "Сторінок ___, рисунків ___, джерел літератури ___, додатків ___.",
    )

    p_en = find_paragraph(doc, "A fullstack application based on Next.js")
    set_text(
        p_en,
        "A fullstack application based on Next.js, MongoDB Atlas, and Auth.js was designed "
        "and implemented. The system supports pet adoption, a searchable and paginated animal "
        "catalog, adoption applications, lost and found pet announcements, user profiles, "
        "administrative functionality, and an AI assistant for personalized pet recommendations.",
    )
    insert_after(
        p_en,
        "The system also includes a separate Animal model for internal animal records, medical "
        "information, documents, found-location details, and information about the person who "
        "found the animal. Cloudinary integration enables image uploads from a computer or "
        "mobile device. The admin area was extended with dashboard analytics, search, filtering, "
        "collapsible lists, and moderation of user-submitted announcements.",
        style=p_en.style.name,
    )
    set_text(
        find_paragraph(doc, "Testing of the system functionality"),
        "Testing of the system functionality, ML module integration, image upload workflow, "
        "user profile functionality, administrative scenarios, announcement moderation, and "
        "containerization process using Docker was carried out.",
    )
    set_text(
        find_paragraph(doc, "The practical significance of the work lies"),
        "The practical significance of the work lies in the creation of a software solution "
        "that can serve as a foundation for a modern digital platform for pet adoption, animal "
        "record management, application processing, lost and found pet announcements, and "
        "administrative decision support through analytics.",
    )
    set_text(
        find_paragraph(doc, "Keywords: AdoptPet"),
        "Keywords: AdoptPet, web application, pet adoption, lost and found pets, admin panel, "
        "analytics, Cloudinary, Auth.js, Next.js, MongoDB, Mongoose, artificial intelligence, "
        "machine learning, computer vision, EfficientNet, FastAPI.",
    )
    set_text(
        find_paragraph(doc, "Pages 106"),
        "Pages ___, figures ___, references ___, appendices ___.",
    )


def update_intro_and_requirements(doc: Document) -> None:
    p = find_paragraph(doc, "Актуальність розробки платформи AdoptPet")
    insert_after(
        p,
        "Після розширення функціональності платформа також вирішує суміжну задачу "
        "централізованого обліку тварин та підтримки сценаріїв, коли тварину знайдено або "
        "загублено. Це важливо для притулків і волонтерів, оскільки одна й та сама тварина "
        "може проходити кілька станів: знайдена, тимчасово утримується, проходить лікування, "
        "очікує адопції або вже прилаштована. Тому система повинна містити не лише публічну "
        "анкету, а й внутрішню картку з медичними, документальними та адміністративними даними.",
        style=p.style.name,
    )
    set_text(
        find_paragraph(doc, "Метою роботи є розробка, дослідження та тестування програмного комплексу AdoptPet"),
        "Метою роботи є розробка, дослідження та тестування програмного комплексу AdoptPet, "
        "який забезпечує цифрову підтримку адопції тварин, ведення внутрішнього обліку тварин, "
        "створення та модерацію оголошень про загублених/знайдених тварин, завантаження "
        "зображень, адміністративну аналітику, AI-рекомендації та автоматичне визначення "
        "породи тварини за зображенням.",
    )
    anchor = find_paragraph(doc, "•\tреалізувати fullstack-застосунок")
    bullet_after(
        anchor,
        [
            "реалізувати окрему модель Animal для збереження службової інформації про тварину;",
            "реалізувати модуль загублених і знайдених тварин із публічним поданням оголошень та адміністративною модерацією;",
            "інтегрувати Cloudinary для завантаження фотографій тварин і фото профілю користувача;",
            "розробити особистий кабінет користувача та dashboard адміністратора з аналітичними показниками;",
        ],
    )
    set_text(
        find_paragraph(doc, "Практичне значення роботи полягає у створенні програмної системи"),
        "Практичне значення роботи полягає у створенні програмної системи, яка може бути "
        "використана притулками, волонтерськими організаціями та користувачами як єдина "
        "платформа для адопції, подання заявок, пошуку загублених/знайдених тварин, "
        "збереження внутрішньої інформації про тварин і аналізу статистики. Розроблена "
        "архітектура передбачає подальше масштабування: додавання мобільного застосунку, "
        "розширення аналітики, інтеграцію платежів, автоматизацію комунікації з притулками "
        "та підключення нових ML-модулів.",
    )
    set_text(
        find_paragraph(doc, "Структура та обсяг роботи."),
        "Структура та обсяг роботи. Робота складається зі вступу, п’ятьох розділів, висновків, "
        "списку використаних джерел та додатків. Кількість сторінок, рисунків і додатків "
        "уточнюється після остаточного оновлення ілюстративних матеріалів та автоматичного "
        "оновлення змісту документа.",
    )

    set_text(
        find_paragraph(doc, "У категорію Must have увійшли вимоги"),
        "У категорію Must have увійшли вимоги, без яких система не може виконувати свою "
        "основну функцію: реєстрація та авторизація користувачів; перегляд каталогу адопції; "
        "пошук, фільтрація та пагінація; профіль тварини; подача заявки; особистий кабінет; "
        "завантаження фото; публічна форма загублених/знайдених тварин; модерація оголошень; "
        "адміністративне керування; базовий AI-помічник.",
    )
    set_text(
        find_paragraph(doc, "До категорії Should have"),
        "До категорії Should have, тобто бажаних аспектів, які покращують досвід користування, "
        "належать: рекомендації на основі AI; повідомлення про статус заявки; історія взаємодії; "
        "dashboard адміністратора; внутрішня картка Animal з медичним станом, документами та "
        "даними про місце знаходження тварини.",
    )
    set_text(
        find_paragraph(doc, "Категорія Could have містить"),
        "Категорія Could have містить функціональність, яка може бути розширена за наявності "
        "додаткового часу: розширена аналітика за регіонами й породами, рейтинг притулків, "
        "відгуки, інтеграція платіжних систем для регулярної допомоги, push-сповіщення та "
        "геолокаційний пошук.",
    )
    functional = find_paragraph(doc, "Функціональні вимоги:")
    # Replace the short old list by appending a fuller list after the heading-like line.
    current = functional
    for item in [
        "реєстрація та авторизація користувача;",
        "перегляд каталогу адопції з пошуком, фільтрацією, пагінацією та клікабельними картками;",
        "перегляд профілю тварини та подання заявки на адопцію;",
        "створення оголошення про загублену або знайдену тварину;",
        "модерація користувацьких оголошень адміністратором;",
        "завантаження фото через Cloudinary;",
        "ведення особистого кабінету користувача та фото профілю;",
        "адміністрування користувачів, оголошень, заявок і карток Animal;",
        "перегляд аналітичного dashboard із кількісними показниками;",
        "AI-рекомендації та ML-визначення породи за зображенням.",
    ]:
        current = insert_after(current, item, style=functional.style.name)
    remove_exact_paragraphs(
        doc,
        {
            "реєстрація користувача;",
            "авторизація;",
            "перегляд каталогу;",
            "AI-рекомендації;",
            "подача заявки;",
            "адміністрування контенту.",
        },
    )


def update_architecture(doc: Document) -> None:
    p = find_paragraph(doc, "Для забезпечення безпеки доступу до системи використовується механізм автентифікації")
    insert_after(
        p,
        "Для роботи із медіафайлами використовується зовнішній сервіс Cloudinary. "
        "Зображення завантажуються через серверний API-маршрут, після чого у MongoDB "
        "зберігається URL оптимізованого зображення. Такий підхід зменшує навантаження "
        "на основний застосунок і дозволяє використовувати однаковий механізм завантаження "
        "для анкет тварин, оголошень про загублених/знайдених тварин і фото профілю користувача.",
        style=p.style.name,
    )
    set_text(
        find_paragraph(doc, "Ключовим компонентом системи є база даних MongoDB"),
        "Ключовим компонентом системи є база даних MongoDB, у якій зберігається основна "
        "предметна інформація. У межах розробленої моделі даних передбачено колекції User, "
        "Post, Animal та AdoptionForm. Колекція Post відповідає за публічні анкети та "
        "оголошення, Animal зберігає розширену внутрішню інформацію про тварину, User містить "
        "облікові дані й фото профілю користувача, а AdoptionForm зберігає заявки на адопцію.",
    )
    set_text(
        find_paragraph(doc, "Основними сутностями системи є користувачі"),
        "Основними сутностями системи є користувачі (User), публічні оголошення та анкети "
        "тварин (Post), службові картки тварин (Animal) і заявки на адопцію (AdoptionForm). "
        "Між сутностями реалізовано логічні зв’язки, які забезпечують взаємодію користувачів "
        "із платформою, синхронізацію публічного оголошення з внутрішньою карткою тварини та "
        "можливість подальшого аналізу даних адміністратором.",
    )
    p = find_paragraph(doc, "Основними сутностями системи є користувачі")
    insert_after(
        p,
        "Модель Animal пов’язана з Post і використовується для збереження даних, які не завжди "
        "мають бути частиною публічної анкети: місце, де тварину знайшли, контактні дані особи, "
        "яка її знайшла, медичний стан, перелік хвороб, документи, внутрішні примітки та "
        "поточний статус перебування. Для вже створених оголошень передбачено backfill/migration, "
        "який створює відповідні записи Animal на основі існуючих Post.",
        style=p.style.name,
    )
    insert_after(
        find_paragraph(doc, "Основними бібліотеками є:"),
        "cloudinary — для завантаження, зберігання та оптимізації зображень;",
    )
    insert_after(
        find_paragraph(doc, "next-auth — для управління сесіями користувачів."),
        "jest та Testing Library — для автоматизованого тестування серверної логіки, API-маршрутів і користувацьких сценаріїв.",
    )
    p = find_paragraph(doc, "Для звичайних користувачів доступ до адміністративного функціоналу обмежений")
    insert_after(
        p,
        "Додатковим елементом безпеки є модерація оголошень про загублених і знайдених тварин. "
        "Користувач може подати інформацію через публічну форму, однак публікація оголошення "
        "відбувається лише після перевірки адміністратором. Завантаження зображень виконується "
        "через серверний маршрут, що дозволяє приховати приватні ключі Cloudinary та контролювати "
        "тип і розмір файлів.",
        style=p.style.name,
    )


def update_implementation(doc: Document) -> None:
    set_text(
        find_paragraph(doc, "Робота з даними реалізована за допомогою бібліотеки Mongoose"),
        "Робота з даними реалізована за допомогою бібліотеки Mongoose, яка виступає як "
        "об’єктно-документна модель (ODM) для MongoDB. У межах проєкту визначено моделі "
        "User, Post, Animal та AdoptionForm. Модель Post відповідає за публічні оголошення "
        "та анкети, Animal — за внутрішній облік тварини, User — за користувачів і профільні "
        "дані, AdoptionForm — за заявки на адопцію.",
    )
    p = find_paragraph(doc, "Робота з даними реалізована за допомогою бібліотеки Mongoose")
    insert_after(
        p,
        "На серверному рівні реалізовано логіку синхронізації Post і Animal. Під час створення "
        "або оновлення оголошення формується відповідна картка тварини, а для старих записів "
        "передбачено backfill-скрипт. Це дозволяє не втратити дані, створені до появи окремої "
        "моделі Animal, і поступово перевести систему до більш повної структури обліку.",
        style=p.style.name,
    )
    insert_after(
        p,
        "Окремий API-маршрут /api/upload відповідає за завантаження зображень у Cloudinary. "
        "Цей маршрут використовується в адміністративних формах, формі загублених/знайдених "
        "тварин і в особистому кабінеті користувача для фото профілю.",
        style=p.style.name,
    )
    p_front = find_paragraph(doc, "Одним із ключових елементів інтерфейсу є навігаційна панель")
    insert_after(
        p_front,
        "Клієнтська частина доповнена окремими сторінками для адопції тварин, загублених і "
        "знайдених тварин, публічної форми створення оголошення, особистого кабінету, разової "
        "допомоги фонду та адміністративного dashboard. У каталозі реалізовано клікабельні "
        "картки, пагінацію та зручний перехід до детальної сторінки тварини.",
        style=p_front.style.name,
    )
    insert_after(
        p_front,
        "Для завантаження зображень створено повторно використовуваний компонент ImageUploadInput. "
        "Він дозволяє користувачам і адміністраторам додавати фото без ручного введення URL, що "
        "покращує зручність роботи із системою та зменшує кількість помилок під час заповнення форм.",
        style=p_front.style.name,
    )
    p_ml = find_paragraph(doc, "Крім того, використання API забезпечує гнучкість інтеграції")
    set_text(
        p_ml,
        "Крім того, використання API забезпечує гнучкість інтеграції та дозволяє підключати "
        "додаткові інтелектуальні функції, зокрема класифікацію порід тварин за зображенням. "
        "Зображення, завантажені через Cloudinary, можуть використовуватися як вхідні дані "
        "для ML-сервісу, а результати прогнозу зберігаються в анкеті тварини у структурованому вигляді.",
    )
    # Fix duplicated wording in chapter 3.
    p_dup = find_paragraph(doc, "На стороні Python-сервісу відбувається обробка отриманих даних.")
    set_text(
        p_dup,
        "На стороні Python-сервісу виконується аналіз отриманих даних із використанням моделей "
        "машинного навчання та алгоритмів рекомендаційної системи. Після виконання обчислень "
        "сервіс формує відповідь у форматі JSON, яка передається назад до Next.js-застосунку.",
    )
    p_conclusion = find_paragraph(doc, "У даному розділі було розглянуто процес реалізації програмного комплексу")
    insert_after(
        p_conclusion,
        "Крім базового функціоналу адопції, у розділі враховано реалізацію модулів загублених/"
        "знайдених тварин, особистого кабінету, Cloudinary-завантаження, внутрішньої картки "
        "Animal та адміністративної аналітики. Це розширює систему від звичайного каталогу до "
        "повноцінної інформаційної платформи для користувачів, волонтерів і адміністратора.",
        style=p_conclusion.style.name,
    )


def update_testing_deployment_conclusion(doc: Document) -> None:
    p = find_paragraph(doc, "Оскільки система реалізована як fullstack-застосунок")
    set_text(
        p,
        "Оскільки система реалізована як fullstack-застосунок на основі Next.js 14, тестування "
        "охоплювало frontend-компоненти React, серверну логіку Next.js API Routes і Server Actions, "
        "механізми автентифікації Auth.js, інтеграцію з інтелектуальним модулем, Cloudinary-"
        "завантаження, модерацію оголошень та адміністративні сценарії роботи.",
    )
    p = find_paragraph(doc, "У процесі тестування перевірялося:")
    current = p
    for item in [
        "коректна робота пагінації каталогу та клікабельних карток;",
        "створення і модерація оголошень про загублених/знайдених тварин;",
        "завантаження фото через Cloudinary у формах адопції, оголошень і профілю;",
        "відображення та редагування картки Animal з медичним станом і документами;",
        "робота dashboard-аналітики за видами тварин, статусами, породами та регіонами;",
    ]:
        current = insert_after(current, item, style=p.style.name)
    p_admin = find_paragraph(doc, "У межах тестування адміністративної частини перевірялися:")
    current = p_admin
    for item in [
        "пошук користувачів за username/email;",
        "пошук оголошень за назвою/slug;",
        "згортання великих адміністративних списків;",
        "перехід з адміністративної панелі до публічної картки тварини;",
        "перегляд графіків і зведених показників у dashboard.",
    ]:
        current = insert_after(current, item, style=p_admin.style.name)
    set_text(
        find_paragraph(doc, "Основний веб-застосунок виконує функції frontend- та backend-рівня"),
        "Основний веб-застосунок виконує функції frontend- та backend-рівня. Він відповідає "
        "за відображення інтерфейсу користувача, маршрутизацію, авторизацію, роботу з каталогом "
        "тварин, обробку заявок на адопцію, модерацію оголошень, особистий кабінет, "
        "адміністративну аналітику, завантаження зображень через Cloudinary та взаємодію з базою даних. "
        "ML-сервіс реалізований окремо на FastAPI та використовується для класифікації породи "
        "тварини за зображенням.",
    )
    duplicate = find_paragraph(doc, "Після запуску користувач взаємодіє з платформою через браузер.")
    second_duplicate = None
    seen = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Після запуску користувач взаємодіє з платформою через браузер."):
            if seen:
                second_duplicate = paragraph
                break
            seen = True
    if second_duplicate is not None:
        remove_paragraph(second_duplicate)
    set_text(
        duplicate,
        "Після запуску користувач взаємодіє з платформою через браузер. Запити до каталогу, "
        "авторизації, особистого кабінету, адміністративної панелі, завантаження зображень, "
        "заявок і оголошень обробляються Next.js-застосунком. Якщо користувач або адміністратор "
        "запускає визначення породи за зображенням, Next.js надсилає HTTP-запит до FastAPI-сервісу, "
        "який формує прогноз і повертає результат у форматі JSON.",
    )
    set_text(
        find_paragraph(doc, "Основною метою створення платформи є спрощення процесу адопції"),
        "Основною метою створення платформи є спрощення процесу адопції тварин, цифровізація "
        "внутрішнього обліку тварин, підтримка оголошень про загублених/знайдених тварин, "
        "автоматизація пошуку інформації про породи та підвищення зручності взаємодії "
        "користувачів із системою.",
    )
    p_eff = find_paragraph(doc, "Крім ML-частини, ефективність рішення також підтверджується")
    insert_after(
        p_eff,
        "Окремим показником ефективності є поява адміністративної аналітики. Dashboard дозволяє "
        "оцінювати кількість собак і котів, статуси тварин, частку прилаштованих тварин, "
        "популярні породи, вікові групи, регіони походження та обсяг заявок. Такі дані можуть "
        "використовуватися адміністратором для планування роботи притулку або волонтерської організації.",
        style=p_eff.style.name,
    )
    set_text(
        find_paragraph(doc, "У результаті виконання дипломної роботи було розроблено"),
        "У результаті виконання дипломної роботи було розроблено та досліджено програмний "
        "комплекс AdoptPet — інтелектуальну веб-платформу для підтримки адопції тварин, "
        "обліку тварин, модерації оголошень, завантаження зображень, роботи особистого "
        "кабінету та аналітичної підтримки адміністратора.",
    )
    set_text(
        find_paragraph(doc, "Практичне значення роботи полягає у створенні програмної системи, яка може використовуватися"),
        "Практичне значення роботи полягає у створенні програмної системи, яка може використовуватися "
        "як основа для цифрової платформи у сфері адопції та волонтерської допомоги тваринам. "
        "Запропоноване рішення дозволяє підвищити ефективність роботи притулків, спростити "
        "процес подання заявок, централізувати інформацію про загублених/знайдених тварин, "
        "зберігати службову інформацію про медичний стан і документи тварин та отримувати "
        "зведену аналітику для прийняття управлінських рішень.",
    )
    set_text(
        find_paragraph(doc, "Таким чином, поставлена у дипломній роботі мета була досягнута"),
        "Таким чином, поставлена у дипломній роботі мета була досягнута, а всі визначені "
        "задачі — виконані. Результатом роботи стала інтелектуальна платформа AdoptPet, яка "
        "поєднує сучасні веб-технології, машинне навчання, AI-сервіси, хмарне зберігання "
        "зображень, модерацію користувацького контенту та адміністративну аналітику для "
        "підтримки процесів адопції й допомоги тваринам.",
    )
    refs_anchor = find_paragraph(doc, "Auth.js | Authentication for the Web.")
    insert_after(
        refs_anchor,
        "Cloudinary Documentation. Image and Video Upload, Storage, Optimization and Delivery. "
        "URL: https://cloudinary.com/documentation (дата звернення: 23.05.2026).",
        style=refs_anchor.style.name,
    )


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    update_annotation(doc)
    update_intro_and_requirements(doc)
    update_architecture(doc)
    update_implementation(doc)
    update_testing_deployment_conclusion(doc)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
