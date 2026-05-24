from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph


SOURCE = Path("outputs/Диплом_оновлена_версія_AdoptPet.docx")
OUTPUT = Path("outputs/Диплом_оновлена_версія_AdoptPet_з_діаграмами.docx")

IMAGES = {
    "context": Path(r"C:\Users\Huawei\OneDrive\Desktop\diagra.png"),
    "component": Path(r"C:\Users\Huawei\OneDrive\Desktop\component.png"),
    "data_model": Path(r"C:\Users\Huawei\OneDrive\Desktop\DataModel.png"),
    "deployment": Path(r"C:\Users\Huawei\OneDrive\Desktop\deployment.png"),
    "ml_pipeline": Path(r"C:\Users\Huawei\OneDrive\Desktop\MLPipeline.png"),
    "sequence_success": Path(r"C:\Users\Huawei\OneDrive\Desktop\sequence1.png"),
    "sequence_error": Path(r"C:\Users\Huawei\OneDrive\Desktop\sequence2.png"),
}


def find_index(doc: Document, startswith: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(startswith):
            return index
    raise ValueError(f"Paragraph not found: {startswith}")


def add_picture_to_paragraph(paragraph: Paragraph, image_path: Path, width_inches: float) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def replace_image_before_caption(doc: Document, caption_prefix: str, image_path: Path, width_inches: float) -> None:
    caption_index = find_index(doc, caption_prefix)
    image_paragraph = doc.paragraphs[caption_index - 1]
    add_picture_to_paragraph(image_paragraph, image_path, width_inches)


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = deepcopy(paragraph._p)
    element.clear_content()
    paragraph._p.addnext(element)
    new_paragraph = Paragraph(element, paragraph._parent)
    if style:
        new_paragraph.style = style
    elif paragraph.style:
        new_paragraph.style = paragraph.style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_before(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = deepcopy(paragraph._p)
    element.clear_content()
    paragraph._p.addprevious(element)
    new_paragraph = Paragraph(element, paragraph._parent)
    if style:
        new_paragraph.style = style
    elif paragraph.style:
        new_paragraph.style = paragraph.style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def main() -> None:
    missing = [str(path) for path in IMAGES.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("\n".join(missing))

    doc = Document(SOURCE)

    replace_image_before_caption(doc, "Рисунок 2.1", IMAGES["context"], 5.8)
    replace_image_before_caption(doc, "Рис. 2.2", IMAGES["component"], 6.45)
    replace_image_before_caption(doc, "Рис. 2.3", IMAGES["data_model"], 6.45)

    replace_image_before_caption(doc, "Рис. 2.7", IMAGES["sequence_success"], 6.25)
    doc.paragraphs[find_index(doc, "Рис. 2.7")].text = (
        "Рис. 2.7. Діаграма послідовності сценарію визначення породи тварини за фото"
    )
    seq_caption = doc.paragraphs[find_index(doc, "Рис. 2.7")]
    seq_error_image = insert_after(seq_caption)
    add_picture_to_paragraph(seq_error_image, IMAGES["sequence_error"], 6.25)
    insert_after(
        seq_error_image,
        "Рис. 2.8. Діаграма послідовності сценарію недоступності ML-сервісу",
        style=seq_caption.style.name,
    )

    doc.paragraphs[find_index(doc, "Рис. 2.8. Контейнерна")].text = (
        "Рис. 2.9. Контейнерна структура системи AdoptPet"
    )
    replace_image_before_caption(doc, "Рис. 2.9", IMAGES["deployment"], 6.45)

    old_ai_caption = doc.paragraphs[find_index(doc, "Рис. 4.13")]
    old_ai_caption.text = "Рис. 4.14. Архітектура AI-модуля платформи AdoptPet"
    heading_46 = doc.paragraphs[find_index(doc, "4.6 Інтеграція AI-помічника")]
    ml_caption = insert_before(
        heading_46,
        "Рис. 4.13. Pipeline навчання, inference та моніторингу ML-модуля",
        style=old_ai_caption.style.name,
    )
    ml_image = insert_before(ml_caption)
    add_picture_to_paragraph(ml_image, IMAGES["ml_pipeline"], 6.45)

    replace_image_before_caption(doc, "Рис. 5.9", IMAGES["deployment"], 6.45)

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
