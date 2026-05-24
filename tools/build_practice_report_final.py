from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\AdoptPet\adoptpetfinal")
OUT = ROOT / "outputs"
SCREENS = OUT / "diploma_ui_screenshots"
DESKTOP = Path(r"C:\Users\Huawei\OneDrive\Desktop")
OUTPUT = OUT / "Звіт_переддипломна_практика_AdoptPet_актуальний.docx"

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
PURPLE = RGBColor(104, 54, 218)


def set_font(run, size=14, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    total = int(sum(widths_cm) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(w_cm / 2.54 * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for cell, w_cm in zip(row.cells, widths_cm):
            cell.width = Cm(w_cm)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def table(doc, headers, rows, widths_cm):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_widths(t, widths_cm)
    for cell, txt in zip(t.rows[0].cells, headers):
        shade_cell(cell, "EEE7FB")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt)
        set_font(r, size=11, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for cell, txt in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(txt)
            set_font(r, size=11)
    doc.add_paragraph()
    return t


def para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, bold=bold, italic=italic)
    return p


def centered(doc, text="", bold=False, size=14, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def heading(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12 if level == 2 else 0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size=14, bold=True)
    return p


def set_alt_text(shape, descr):
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", descr)


def figure(doc, img_path, caption, width_cm=15.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(img_path), width=Cm(width_cm))
    set_alt_text(shape, caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    r = cap.add_run(caption)
    set_font(r, size=12)


def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=12)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=12)


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Зміст оновлюється автоматично у Microsoft Word (Ctrl+A, F9)."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])
    set_font(run, size=14)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(14)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name in ("Heading 1", "Heading 2"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0)

    page_number(section.footer.paragraphs[0])
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def cover(doc):
    centered(doc, "Міністерство освіти і науки України", bold=False)
    centered(doc, "Чернівецький національний університет імені Юрія Федьковича")
    centered(doc, "Навчально-науковий інститут фізико-технічних та комп'ютерних наук")
    centered(doc, "Кафедра комп'ютерних наук")
    for _ in range(4):
        doc.add_paragraph()
    centered(doc, "ЗВІТ", bold=True, size=16)
    centered(doc, "з переддипломної практики", bold=True)
    centered(doc, "на тему:", bold=False)
    centered(
        doc,
        "«Розробка інтелектуальної веб-платформи AdoptPet для підтримки адопції тварин»",
        bold=True,
    )
    for _ in range(3):
        doc.add_paragraph()
    para(doc, "Студентки 4-го курсу, групи 444а", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    para(doc, "спеціальності 122 «Комп'ютерні науки»", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    para(doc, "Сидорець Аміни Олександрівни", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    para(doc, "База практики: кафедра комп'ютерних наук", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    para(doc, "Керівник від кафедри: ____________________", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    para(doc, "Термін практики: ___________________________", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    for _ in range(4):
        doc.add_paragraph()
    centered(doc, "Чернівці - 2026")


def build():
    doc = Document()
    configure_document(doc)
    cover(doc)

    heading(doc, "РЕФЕРАТ", page_break=True)
    para(
        doc,
        "Звіт присвячено результатам проходження переддипломної практики, у межах якої "
        "розроблено та апробовано програмний комплекс AdoptPet - fullstack-вебплатформу "
        "для підтримки адопції тварин, ведення внутрішнього обліку, приймання заявок, "
        "модерації оголошень про загублених і знайдених тварин та адміністративної аналітики."
    )
    para(
        doc,
        "Під час практики реалізовано каталог адопції з пошуком, фільтрами та пагінацією, "
        "детальні картки тварин, форму заявки, особистий кабінет користувача, публічну "
        "форму lost/found із модерацією, окрему модель Animal, завантаження зображень через "
        "Cloudinary і dashboard для адміністратора. Інтелектуальна складова охоплює "
        "AI-помічника на основі LangChain/OpenAI та ML-сервіс визначення породи за фотографією."
    )
    para(
        doc,
        "Програмну реалізацію побудовано на основі Next.js 14, React, MongoDB Atlas, "
        "Mongoose та Auth.js. ML-сервіс реалізовано окремим FastAPI-застосунком із моделями "
        "EfficientNet-B3, підготовленими у PyTorch Lightning. Перевірено API-маршрути, "
        "синхронізацію Post і Animal, аналітичні обчислення, взаємодію з AI та обробку "
        "недоступності ML-сервісу."
    )
    para(
        doc,
        "Сторінок - ___, рисунків - 19, таблиць - 5, використаних джерел - 15, додатків - 1.",
        indent=False,
    )
    para(
        doc,
        "Ключові слова: AdoptPet, переддипломна практика, адопція тварин, Next.js, MongoDB, "
        "Animal, Cloudinary, Auth.js, FastAPI, EfficientNet-B3, машинне навчання, "
        "AI-помічник, адміністративна аналітика.",
        indent=False,
    )

    heading(doc, "ЗМІСТ", page_break=True)
    add_toc_field(doc)

    heading(doc, "ВСТУП", page_break=True)
    para(
        doc,
        "Переддипломна практика є завершальним етапом професійної підготовки бакалавра, "
        "під час якого результати навчання застосовуються у повному циклі створення "
        "програмного продукту: від аналізу предметної області та формування вимог до "
        "реалізації, тестування і документування. Індивідуальним завданням практики стала "
        "розробка веб-платформи AdoptPet, орієнтованої на допомогу тваринам та організацію "
        "процесу адопції."
    )
    para(
        doc,
        "Актуальність теми пов'язана з тим, що оголошення про тварин часто поширюються "
        "через розрізнені канали: соціальні мережі, месенджери та локальні списки. У такому "
        "середовищі складно контролювати актуальність даних, статус тварини, заявки "
        "користувачів, історію лікування та інформацію про місце знахідки. Єдина система "
        "дозволяє структурувати ці процеси й забезпечити керований доступ до інформації."
    )
    para(
        doc,
        "Метою практики було удосконалення та завершення функціональної версії AdoptPet "
        "із підтримкою публічних і адміністративних сценаріїв, роботи із зображеннями, "
        "аналітики й інтелектуальних компонентів. Об'єктом роботи є процес створення "
        "інформаційної системи підтримки адопції тварин, а предметом - архітектурні та "
        "програмні засоби її реалізації з використанням сучасних веб- та ML-технологій."
    )
    para(
        doc,
        "У звіті наведено аналіз завдання, характеристику створених модулів, опис "
        "архітектури й моделі даних, результати реалізації Cloudinary та ML-інтеграції, "
        "перевірку функціональності та оцінку результатів практики."
    )

    heading(doc, "1 АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ ТА ПОСТАНОВКА ІНДИВІДУАЛЬНОГО ЗАВДАННЯ", page_break=True)
    heading(doc, "1.1 Характеристика предметної області", level=2)
    para(
        doc,
        "Предметна область платформи охоплює публікацію інформації про тварин, що "
        "потребують дому, подання заявок потенційними власниками, пошук загублених і "
        "знайдених тварин, а також внутрішній облік. У межах такої системи тварина не є "
        "лише фотографією та коротким описом: для відповідального рішення потрібні дані "
        "про статус, медичні особливості, документи, місцезнаходження та історію звернень."
    )
    para(
        doc,
        "Користувач системи повинен мати можливість знайти тварину за зрозумілими "
        "параметрами, переглянути її детальну картку та подати заявку. Для випадків "
        "зникнення або знахідки потрібна окрема оперативна форма оголошення. Адміністратор "
        "відповідає за перевірку контенту, ведення службових даних і перегляд аналітики, "
        "яка допомагає оцінити кількість тварин і результативність процесів."
    )
    figure(doc, DESKTOP / "ContextDiagram.png", "Рисунок 1.1 - Контекстна діаграма системи AdoptPet", 15.2)
    para(
        doc,
        "На контекстній діаграмі показано фактичні ролі користувача й адміністратора та "
        "зовнішні сервіси, із якими взаємодіє система. Роль волонтера або притулку "
        "позначено як перспективну, оскільки у поточній версії окремий кабінет цієї ролі "
        "не реалізовано."
    )
    heading(doc, "1.2 Актуальність та проблеми, що вирішує платформа", level=2)
    para(
        doc,
        "Проблемою наявних неформалізованих каналів є відсутність єдиної моделі даних. "
        "Навіть якщо оголошення має фотографію, у ньому може не бути інформації про "
        "статус тварини, місто, контакт, медичний стан або результат розгляду заявки. "
        "Повторне дублювання публікацій ускладнює пошук і підвищує ризик використання "
        "неактуальної інформації."
    )
    para(
        doc,
        "AdoptPet орієнтовано на вирішення цієї проблеми шляхом об'єднання публічної "
        "частини, особистого кабінету й адміністративного керування. Система підтримує "
        "адопцію, lost/found-оголошення з модерацією, збереження зображень, внутрішню "
        "картку Animal та dashboard із зведеними показниками. Інтелектуальні модулі "
        "використовуються як допоміжний інструмент, а не як заміна рішення адміністратора."
    )
    heading(doc, "1.3 Мета та завдання переддипломної практики", level=2)
    para(
        doc,
        "Метою індивідуального завдання є розробка і перевірка програмного комплексу "
        "AdoptPet, який забезпечує цифрову підтримку адопції тварин, внутрішнього обліку, "
        "модерації оголошень та прийняття адміністративних рішень на основі даних."
    )
    table_caption(doc, "Таблиця 1.1 - Основні завдання переддипломної практики")
    table(
        doc,
        ["№", "Завдання практики", "Практичний результат"],
        [
            ("1", "Проаналізувати предметну область і сформувати вимоги", "Уточнено ролі, сценарії та межі системи"),
            ("2", "Реалізувати користувацькі сценарії адопції і lost/found", "Каталог, картка, заявка, форма оголошення і модерація"),
            ("3", "Розширити структуру обліку тварин", "Модель Animal і backfill для наявних Post"),
            ("4", "Підключити роботу із зображеннями та інтелектуальні сервіси", "Cloudinary, AI-помічник і FastAPI ML-сервіс"),
            ("5", "Створити адміністративні інструменти й аналітику", "Пошук, згортані списки, редагування Animal, dashboard"),
            ("6", "Перевірити працездатність рішення", "Автоматизовані тести й оцінка сценаріїв"),
        ],
        [0.8, 7.0, 7.2],
    )
    heading(doc, "1.4 Функціональні та нефункціональні вимоги", level=2)
    para(
        doc,
        "До основних функціональних вимог належать реєстрація та вхід користувача, "
        "перегляд каталогу адопції, пошук, фільтрація і сторінкова навігація, детальна "
        "картка тварини, подання заявки, створення оголошення про зникнення або знахідку, "
        "завантаження фотографії та перегляд власних даних у кабінеті."
    )
    para(
        doc,
        "Адміністративні вимоги включають роботу з користувачами, постами, заявками та "
        "Animal-картками; зміну статусів; модерацію lost/found; пошук записів; згортання "
        "великих списків; доступ до dashboard. Аналітичний модуль має відображати кількість "
        "котів і собак, статуси адопції, розподіли за породами, віком і регіонами."
    )
    para(
        doc,
        "Нефункціональні вимоги стосуються захисту адміністративних маршрутів, безпечної "
        "роботи із ключами зовнішніх сервісів, валідації файлів, адаптивного інтерфейсу, "
        "керованої обробки помилок і можливості контейнерного запуску веб-застосунку та "
        "ML-сервісу."
    )

    heading(doc, "2 ПРОЄКТУВАННЯ ТА РЕАЛІЗАЦІЯ ПРОГРАМНОГО КОМПЛЕКСУ ADOPTPET", page_break=True)
    heading(doc, "2.1 Архітектура та технологічний стек", level=2)
    para(
        doc,
        "Платформу реалізовано як fullstack-застосунок на Next.js 14 і React. Next.js "
        "відповідає за серверне формування сторінок, API Routes і Server Actions, тому "
        "користувацькі сценарії та серверна бізнес-логіка підтримуються в єдиному "
        "проєкті. Для стилізації інтерфейсу застосовано CSS Modules."
    )
    para(
        doc,
        "Постійні дані зберігаються у MongoDB Atlas через Mongoose. Auth.js керує "
        "автентифікацією та розмежуванням користувацького й адміністративного доступу. "
        "Зображення передаються на Cloudinary серверним маршрутом `/api/upload`, а "
        "класифікація породи виконується окремим сервісом FastAPI через маршрут "
        "`/api/predict`. Для текстових AI-функцій використано LangChain та OpenAI."
    )
    table_caption(doc, "Таблиця 2.1 - Технологічний стек платформи")
    table(
        doc,
        ["Рівень", "Технології", "Призначення"],
        [
            ("Інтерфейс", "Next.js 14, React, CSS Modules", "Сторінки, форми, компоненти та навігація"),
            ("Серверна логіка", "API Routes, Server Actions", "Валідація, модерація, інтеграції"),
            ("Дані", "MongoDB Atlas, Mongoose", "Зберігання моделей і вибірки"),
            ("Безпека", "Auth.js", "Сесії та доступ адміністратора"),
            ("Медіа", "Cloudinary", "Фото тварин і профілю"),
            ("AI / ML", "OpenAI, LangChain, FastAPI, PyTorch", "Рекомендації та класифікація породи"),
        ],
        [3.0, 5.1, 6.9],
    )
    figure(doc, DESKTOP / "ComponentDiagram.png", "Рисунок 2.1 - Компонентна структура платформи AdoptPet", 15.3)
    para(
        doc,
        "Компонентна структура відображає фактичне розділення модулів: каталог, "
        "lost/found, заявки, профіль, керування Animal, модерацію та dashboard реалізовано "
        "на веб-рівні, тоді як зображення й класифікація обробляються через окремі "
        "інтеграційні інтерфейси."
    )
    heading(doc, "2.2 Модель даних і синхронізація картки Animal", level=2)
    para(
        doc,
        "Модель даних платформи базується на чотирьох основних сутностях: User, Post, "
        "Animal та AdoptionForm. User містить реєстраційні дані, фото профілю і ознаку "
        "адміністратора. Post є публічною анкетою або оголошенням і зберігає назву, "
        "зображення, тип оголошення, статус модерації, характеристики тварини та "
        "ML-прогнози."
    )
    para(
        doc,
        "Animal створюється у зв'язку один-до-одного з Post за полем `postId` типу "
        "ObjectId і використовується для службового обліку. Поточна реалізація містить "
        "поля `foundLocation`, `foundByName`, `foundByContact`, масив `diseases` і масив "
        "документів формату `{name, url}`. AdoptionForm зберігає контакти заявника, його "
        "досвід, повідомлення та статус `pending`, `approved` або `rejected`."
    )
    para(
        doc,
        "Щоб окрема картка Animal могла використовуватися для вже наявних оголошень, "
        "у проєкті підготовлено сценарій `scripts/backfill-animals.js`. Він створює "
        "відсутні пов'язані записи з початковими значеннями полів і не потребує видалення "
        "старих постів."
    )
    figure(doc, DESKTOP / "DataModelDiagram.png", "Рисунок 2.2 - Модель даних AdoptPet", 15.2)
    table_caption(doc, "Таблиця 2.2 - Призначення основних моделей даних")
    table(
        doc,
        ["Модель", "Ключові поля фактичної реалізації", "Роль у системі"],
        [
            ("User", "username, email, img, isAdmin", "Обліковий запис і права"),
            ("Post", "listingType, moderationStatus, breed, mlPredictions, status", "Публічна картка або оголошення"),
            ("Animal", "postId, foundLocation, diseases, documents", "Внутрішній облік тварини"),
            ("AdoptionForm", "postId, contacts, experience, status", "Опрацювання заявки"),
        ],
        [2.8, 7.1, 5.1],
    )
    heading(doc, "2.3 Користувацькі сценарії", level=2)
    para(
        doc,
        "Основною публічною функцією системи є розділ «Адопція тварин». Каталог "
        "відображає лише оголошення адопції, підтримує пошук, фільтри за типом, віком і "
        "розміром, а також пагінацію. API каталогу повертає по дев'ять елементів на "
        "сторінку за замовчуванням і обмежує максимальний розмір сторінки до 24 записів."
    )
    figure(doc, SCREENS / "02_catalog.png", "Рисунок 2.3 - Каталог адопції з фільтрацією та картками тварин", 15.2)
    para(
        doc,
        "Картка у каталозі є клікабельною і веде до детальної сторінки. На ній "
        "користувач переглядає основні характеристики, AI-опис породи, статус доступності "
        "та може розгорнути відомості про медичний стан. Для доступної тварини відкривається "
        "форма заявки на адопцію."
    )
    figure(doc, SCREENS / "10_animal_details.png", "Рисунок 2.4 - Детальна картка тварини з блоком медичного стану", 14.7)
    para(
        doc,
        "Другим публічним сценарієм є розділ «Загублені / знайдені». Авторизований "
        "користувач може подати інформацію про зниклу або знайдену тварину, вказати місто, "
        "прикмети, місце знахідки та додати фото. Нове оголошення отримує статус "
        "`pending` і стає загальнодоступним після модерації адміністратором."
    )
    figure(doc, SCREENS / "03_lost_found.png", "Рисунок 2.5 - Розділ загублених і знайдених тварин", 15.2)
    figure(doc, SCREENS / "09_lost_found_form.png", "Рисунок 2.6 - Публічна форма створення lost/found-оголошення", 14.2)
    para(
        doc,
        "Особистий кабінет дозволяє користувачеві переглядати власний профіль, створені "
        "оголошення та заявки. Зображення профілю можна вибрати на комп'ютері чи телефоні; "
        "за необхідності збережено й варіант введення URL. Сторінка разової допомоги "
        "містить підготовлений інтерфейс вибору внеску, але інтеграція фактичного "
        "платіжного провайдера належить до подальшого розвитку."
    )
    heading(doc, "2.4 Адміністративна панель та аналітика", level=2)
    para(
        doc,
        "Адміністративна панель є робочим середовищем для опрацювання даних системи. "
        "Реалізовано керування постами й заявками, пошук користувачів за username або "
        "email, пошук оголошень за назвою або slug, зміну статусів та згортані списки. "
        "В окремому розділі «Тварини» адміністратор може редагувати поля Animal і "
        "переходити до публічної картки відповідної тварини."
    )
    figure(doc, SCREENS / "12_admin_panel_top.png", "Рисунок 2.7 - Адміністративна панель керування платформою", 15.2)
    para(
        doc,
        "Dashboard виконує зведення поточних даних. Обчислення формуються на основі "
        "постів, форм адопції та користувачів: загальна кількість записів тварин, "
        "кількість доступних, зарезервованих і вже адоптованих тварин, lost/found-записи, "
        "заявки, розподіли за типом, породою, віком, статусом і містом."
    )
    figure(doc, SCREENS / "06_dashboard.png", "Рисунок 2.8 - Dashboard зі зведеною аналітикою тварин", 15.2)
    heading(doc, "2.5 Завантаження зображень і інтелектуальні модулі", level=2)
    para(
        doc,
        "Для роботи з фотографіями реалізовано серверний маршрут `POST /api/upload`. "
        "Маршрут перевіряє наявність файла, його MIME-тип і обмеження розміру до 5 МБ, "
        "після чого формує підписаний запит до Cloudinary. Клієнт отримує `secure_url` і "
        "`public_id`, а в базі даних зберігається посилання. Приватні ключі Cloudinary "
        "залишаються лише у серверному середовищі."
    )
    para(
        doc,
        "AI-помічник працює з текстовим запитом користувача та актуальними доступними "
        "постами, формуючи рекомендації природною мовою. У програмній реалізації "
        "використано LangChain та модель `gpt-4o-mini`. Окремо реалізовано генерацію "
        "опису породи та набору характеристик для картки тварини."
    )
    para(
        doc,
        "ML-компонент призначений для визначення породи за фото. Сервіс FastAPI приймає "
        "`imageUrl` і тип тварини, вибирає checkpoint для котів або собак, завантажує "
        "зображення, приводить його до розміру 300x300, нормалізує та передає до "
        "EfficientNet-B3. У відповідь повертаються `bestPrediction` і п'ять "
        "`topPredictions` із рівнем впевненості."
    )
    figure(doc, DESKTOP / "MLPipeline.png", "Рисунок 2.9 - Pipeline навчання та використання ML-модуля", 15.1)
    para(
        doc,
        "Моделі навчалися на датасетах Oxford-IIIT Pet, Stanford Dogs і Kaggle Dog "
        "Breed Identification. У підключених картках класів Oxford містить 35 класів, "
        "а Stanford і Kaggle - по 120 порід. За матеріалами дослідження найвище значення "
        "validation Top-1 отримано для Oxford-IIIT Pet - близько 0,93; для Stanford Dogs "
        "результат перевищує 0,87, для Kaggle становить приблизно 0,85."
    )
    para(
        doc,
        "У наскрізному сценарії адміністратор завантажує фото через Cloudinary, після "
        "чого Next.js-маршрут `/api/predict` проксіює запит до FastAPI. Прогноз може бути "
        "збережено у Post разом із пов'язаною карткою Animal. Якщо ML-сервіс тимчасово "
        "недоступний, система повертає контрольовану помилку і дозволяє створити пост "
        "без автоматичного визначення породи."
    )
    table_caption(doc, "Таблиця 2.3 - Серверні інтерфейси реалізованих інтеграцій")
    table(
        doc,
        ["Маршрут", "Вхідні дані", "Результат"],
        [
            ("POST /api/upload", "Файл зображення", "Cloudinary secure URL та publicId"),
            ("POST /api/predict", "imageUrl, type", "bestPrediction і topPredictions"),
            ("GET /api/catalog", "search, type, age, size, page", "Сторінка публічних анкет"),
            ("POST /api/agent/chat", "Текстовий запит", "Рекомендація AI-помічника"),
        ],
        [3.5, 5.5, 6.0],
    )
    figure(doc, DESKTOP / "SequenceML1.png", "Рисунок 2.10 - Успішний сценарій визначення породи при створенні поста", 15.4)
    figure(doc, DESKTOP / "SequenceML2.png", "Рисунок 2.11 - Обробка недоступності ML-сервісу", 15.4)

    heading(doc, "3 ТЕСТУВАННЯ, АПРОБАЦІЯ ТА РЕЗУЛЬТАТИ ВИКОНАННЯ ПРАКТИКИ", page_break=True)
    heading(doc, "3.1 Перевірка реалізованої функціональності", level=2)
    para(
        doc,
        "Перевірка виконувалася для основних користувацьких, адміністративних та "
        "інтеграційних сценаріїв. До автоматизованих тестів увійшли маршрути каталогу, "
        "фільтрація в адміністративній частині, розрахунок dashboard-аналітики, "
        "синхронізація Post і Animal, AI API та новий проксі-маршрут `/api/predict`."
    )
    para(
        doc,
        "Під час фінальної перевірки успішно виконано сім наборів Jest-тестів, що "
        "містять 22 тести, а також статичну перевірку ESLint. Для маршруту прогнозування "
        "перевірено позитивну відповідь, відхилення запиту без обов'язкових даних і "
        "керовану відповідь зі статусом 503, коли ML-сервіс недоступний."
    )
    table_caption(doc, "Таблиця 3.1 - Перевірка основних сценаріїв роботи системи")
    table(
        doc,
        ["Сценарій", "Що перевірено", "Результат"],
        [
            ("Каталог", "Пошук, фільтри, виключення lost/found, пагінація", "Виконано"),
            ("Animal", "Створення та оновлення пов'язаної картки", "Виконано"),
            ("Lost/found", "Створення post зі статусом pending і модерація", "Виконано"),
            ("Dashboard", "Підрахунки статусів, типів, заявок і користувачів", "Виконано"),
            ("AI API", "Валідна відповідь і обробка помилки", "Виконано"),
            ("ML proxy", "200, 400 та 503 для `/api/predict`", "Виконано"),
        ],
        [3.0, 9.0, 3.0],
    )
    heading(doc, "3.2 Розгортання та інтеграційна апробація", level=2)
    para(
        doc,
        "Схема розгортання передбачає запуск веб-застосунку та ML-сервісу як окремих "
        "компонентів. У файлі `docker-compose.yml` для сервісу `web` налаштовано змінну "
        "`ML_SERVICE_URL=http://ml-service:8000`, що забезпечує взаємодію контейнерів. "
        "MongoDB Atlas, Cloudinary й OpenAI використовуються як зовнішні хмарні сервіси."
    )
    figure(doc, DESKTOP / "DeploymentDiagram.png", "Рисунок 3.1 - Діаграма розгортання AdoptPet", 15.2)
    para(
        doc,
        "Апробація користувацького інтерфейсу показала, що в межах єдиного стилю "
        "працюють головна сторінка, каталог, lost/found, допомога, профіль і "
        "адміністративний розділ. Відокремлення storage і ML-сервісу дозволяє змінювати "
        "інтелектуальну модель або спосіб зберігання фото без переробки основного "
        "інтерфейсу."
    )
    heading(doc, "3.3 Отримані результати та напрями розвитку", level=2)
    para(
        doc,
        "У результаті виконано повний функціональний цикл: користувач може створити "
        "акаунт, знайти тварину, переглянути деталі й подати заявку; окремо може "
        "повідомити про загублену або знайдену тварину з фотографією. Адміністратор "
        "керує записами, модерує оголошення, редагує внутрішні картки й аналізує "
        "зведені показники."
    )
    para(
        doc,
        "Подальший розвиток доцільно спрямувати на підключення платіжного провайдера "
        "для сторінки разової допомоги, реалізацію окремої ролі притулку або волонтера, "
        "географічний пошук lost/found, сповіщення користувачів, а також виробничий "
        "моніторинг ML-моделей і відстеження зміни якості прогнозів."
    )

    heading(doc, "ВИСНОВКИ", page_break=True)
    para(
        doc,
        "Під час переддипломної практики виконано комплекс робіт із розробки, "
        "розширення та перевірки веб-платформи AdoptPet. Проведено аналіз предметної "
        "області, сформовано вимоги, уточнено архітектуру, розвинено модель даних та "
        "створено функціональні модулі для користувача й адміністратора."
    )
    para(
        doc,
        "Основним результатом стало створення цілісної інформаційної системи, а не "
        "лише каталогу. У платформі реалізовано публічні сценарії адопції і lost/found, "
        "особистий кабінет, Cloudinary-завантаження, окрему картку Animal з даними про "
        "хвороби, документи і знахідку, а також dashboard для внутрішньої аналітики."
    )
    para(
        doc,
        "Застосування AI-помічника та ML-модуля розширює можливості роботи з інформацією. "
        "Окремий FastAPI-сервіс повертає прогноз породи за фото, а серверний проксі "
        "забезпечує керовану інтеграцію з веб-застосунком. Результати тестування "
        "підтвердили працездатність реалізованих сценаріїв та готовність системи до "
        "подальшого розвитку в межах кваліфікаційної роботи."
    )

    heading(doc, "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", page_break=True)
    sources = [
        "Next.js Documentation. URL: https://nextjs.org/docs (дата звернення: 24.05.2026).",
        "React Documentation. URL: https://react.dev/ (дата звернення: 24.05.2026).",
        "Auth.js Documentation. URL: https://authjs.dev/ (дата звернення: 24.05.2026).",
        "MongoDB Atlas Documentation. URL: https://www.mongodb.com/docs/atlas/ (дата звернення: 24.05.2026).",
        "Mongoose Documentation. URL: https://mongoosejs.com/docs/ (дата звернення: 24.05.2026).",
        "Cloudinary Documentation. Image Upload API Reference. URL: https://cloudinary.com/documentation (дата звернення: 24.05.2026).",
        "FastAPI Documentation. URL: https://fastapi.tiangolo.com/ (дата звернення: 24.05.2026).",
        "PyTorch Documentation. URL: https://pytorch.org/docs/ (дата звернення: 24.05.2026).",
        "PyTorch Lightning Documentation. URL: https://lightning.ai/docs/pytorch/ (дата звернення: 24.05.2026).",
        "LangChain Documentation. URL: https://python.langchain.com/docs/ (дата звернення: 24.05.2026).",
        "OpenAI API Documentation. URL: https://platform.openai.com/docs (дата звернення: 24.05.2026).",
        "Tan M., Le Q. V. EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks. Proceedings of ICML, 2019.",
        "Khosla A. et al. Novel Dataset for Fine-Grained Image Categorization: Stanford Dogs. CVPR Workshop, 2011.",
        "Parkhi O. M. et al. Cats and Dogs: Oxford-IIIT Pet Dataset. CVPR, 2012.",
        "Kaggle. Dog Breed Identification. URL: https://www.kaggle.com/c/dog-breed-identification (дата звернення: 24.05.2026).",
    ]
    for i, source in enumerate(sources, 1):
        para(doc, f"{i}. {source}", indent=False)

    heading(doc, "ДОДАТОК А", page_break=True)
    centered(doc, "ГРАФІЧНІ МАТЕРІАЛИ РЕАЛІЗОВАНОЇ ПЛАТФОРМИ", bold=True)
    para(
        doc,
        "У додатку наведено актуальні екранні форми програмного продукту, які "
        "підтверджують виконання індивідуального завдання під час практики."
    )
    figure(doc, SCREENS / "01_home.png", "Рисунок А.1 - Головна сторінка AdoptPet", 15.2)
    figure(doc, SCREENS / "05_profile.png", "Рисунок А.2 - Особистий кабінет із завантаженням фото профілю", 15.0)
    figure(doc, SCREENS / "11_adoption_form.png", "Рисунок А.3 - Форма подання заявки на адопцію", 13.5)
    figure(doc, SCREENS / "13_admin_panel_lists.png", "Рисунок А.4 - Опрацювання заявок і карток тварин в адмінпанелі", 15.0)
    figure(doc, SCREENS / "04_donation.png", "Рисунок А.5 - Інтерфейс сторінки разової допомоги тваринам", 15.0)
    figure(doc, SCREENS / "08_ai_assistant.png", "Рисунок А.6 - Діалогове вікно AI-помічника", 6.0)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build()
