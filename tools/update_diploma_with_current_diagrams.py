from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph


SOURCE = Path(
    r"C:\Users\Huawei\OneDrive\Desktop\Диплом_оновлена_версія_AdoptPet_з_діаграмами.docx"
)
OUTPUT = Path("outputs/Диплом_AdoptPet_актуалізована_з_новими_діаграмами.docx")

IMAGES = {
    "context": Path(r"C:\Users\Huawei\OneDrive\Desktop\ContextDiagram.png"),
    "component": Path(r"C:\Users\Huawei\OneDrive\Desktop\ComponentDiagram.png"),
    "data_model": Path(r"C:\Users\Huawei\OneDrive\Desktop\DataModelDiagram.png"),
    "deployment": Path(r"C:\Users\Huawei\OneDrive\Desktop\DeploymentDiagram.png"),
    "ml_pipeline": Path(r"C:\Users\Huawei\OneDrive\Desktop\MLPipeline.png"),
    "sequence_success": Path(r"C:\Users\Huawei\OneDrive\Desktop\SequenceML1.png"),
    "sequence_error": Path(r"C:\Users\Huawei\OneDrive\Desktop\SequenceML2.png"),
}


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def find_index(doc: Document, prefix: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index
    raise ValueError(f"Paragraph not found: {prefix}")


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    element = deepcopy(paragraph._p)
    element.clear_content()
    paragraph._p.addnext(element)
    result = Paragraph(element, paragraph._parent)
    if style:
        result.style = style
    elif paragraph.style:
        result.style = paragraph.style
    result.add_run(text)
    return result


def add_picture_to_paragraph(
    paragraph: Paragraph, image_path: Path, width_inches: float
) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def replace_image_before_caption(
    doc: Document, caption_prefix: str, image_path: Path, width_inches: float
) -> Paragraph:
    caption_index = find_index(doc, caption_prefix)
    image_paragraph = doc.paragraphs[caption_index - 1]
    add_picture_to_paragraph(image_paragraph, image_path, width_inches)
    return doc.paragraphs[caption_index]


def replace_diagrams(doc: Document) -> None:
    caption = replace_image_before_caption(doc, "Рисунок 2.1", IMAGES["context"], 6.2)
    set_text(caption, "Рисунок 2.1. Контекстна діаграма системи AdoptPet")

    caption = replace_image_before_caption(doc, "Рис. 2.2", IMAGES["component"], 6.45)
    set_text(caption, "Рис. 2.2. Компонентна діаграма системи AdoptPet")

    caption = replace_image_before_caption(doc, "Рис. 2.3", IMAGES["data_model"], 6.45)
    set_text(caption, "Рис. 2.3. Модель даних платформи AdoptPet")

    caption = replace_image_before_caption(
        doc, "Рис. 2.7", IMAGES["sequence_success"], 6.45
    )
    set_text(
        caption,
        "Рис. 2.7. Діаграма послідовності успішного визначення породи при створенні поста",
    )

    caption = replace_image_before_caption(
        doc, "Рис. 2.8", IMAGES["sequence_error"], 6.45
    )
    set_text(
        caption,
        "Рис. 2.8. Діаграма послідовності сценарію недоступності ML-сервісу",
    )

    caption = replace_image_before_caption(
        doc, "Рис. 2.9", IMAGES["deployment"], 6.45
    )
    set_text(caption, "Рис. 2.9. Діаграма розгортання системи AdoptPet")

    caption = replace_image_before_caption(
        doc, "Рис. 4.13", IMAGES["ml_pipeline"], 6.45
    )
    set_text(
        caption,
        "Рис. 4.13. Pipeline навчання та inference ML-модуля платформи AdoptPet",
    )

    caption = replace_image_before_caption(
        doc, "Рис. 5.9", IMAGES["deployment"], 6.45
    )
    set_text(caption, "Рис. 5.9. Діаграма розгортання системи AdoptPet")


def update_architecture_text(doc: Document) -> None:
    set_text(
        find_paragraph(doc, "Даний підхід дозволяє поєднати клієнтську"),
        "Контекстна діаграма на рисунку 2.1 відображає межі реалізованої системи "
        "та її основні взаємодії. Користувач переглядає каталог адопції, подає "
        "заявки, створює оголошення про загублену або знайдену тварину, "
        "завантажує фотографії й керує профілем. Адміністратор виконує "
        "модерацію, редагує внутрішні картки Animal та переглядає dashboard. "
        "Взаємодія волонтера або притулку на схемі позначена як перспективний "
        "напрям розширення, а не як окрема реалізована роль.",
    )
    set_text(
        find_paragraph(doc, "Окремим компонентом є зовнішні API"),
        "До зовнішніх і спеціалізованих сервісів поточної архітектури належать "
        "Cloudinary для зберігання та обробки зображень, OpenAI/LangChain для "
        "AI-помічника й генерації описів, окремий FastAPI ML-сервіс для "
        "визначення породи за фотографією та Auth.js для автентифікації. "
        "Зображення передаються до Cloudinary лише через серверний маршрут "
        "/api/upload, а ML-запити проходять через маршрут Next.js /api/predict.",
    )
    insert_after(
        find_paragraph(doc, "Рис. 2.3."),
        "Діаграма подає логічний склад основних даних. У фактичній реалізації "
        "MongoDB і Mongoose використовують ObjectId для зв'язків Post, Animal "
        "і AdoptionForm; поля ageGroups та sizes у Post містять одне обране "
        "значення типу String; поле documents у Animal є масивом об'єктів з "
        "назвою документа та URL. Таке уточнення узгоджує логічну схему з "
        "реалізованими моделями застосунку.",
    )
    insert_after(
        find_paragraph(doc, "Рис. 2.8."),
        "На рисунках 2.7 і 2.8 показано фактичний сценарій роботи ML-функції "
        "під час створення поста адміністратором. Фото спочатку завантажується "
        "через /api/upload у Cloudinary, а отриманий imageUrl передається до "
        "серверного маршруту /api/predict, який проксіює запит до FastAPI. "
        "Успішний прогноз зберігається в Post разом із пов'язаною карткою "
        "Animal. Якщо ML-сервіс недоступний, інтерфейс показує повідомлення "
        "про помилку та дозволяє створити пост без автоматично визначеної породи.",
    )
    set_text(
        find_paragraph(doc, "У межах даної системи застосунок логічно"),
        "У поточній реалізації контейнерне середовище складається з двох "
        "сервісів: web, який запускає fullstack-застосунок Next.js, та "
        "ml-service, який запускає FastAPI/PyTorch-сервіс визначення породи.",
    )
    set_text(
        find_paragraph(doc, "Frontend/Backend (Next.js fullstack)"),
        "Сервіс web відповідає за інтерфейс, API Routes, Server Actions, "
        "автентифікацію, роботу з MongoDB Atlas, Cloudinary та OpenAI API.",
    )
    set_text(
        find_paragraph(doc, "База даних (MongoDB Atlas)"),
        "MongoDB Atlas, Cloudinary, OpenAI API та провайдер автентифікації "
        "використовуються як зовнішні хмарні сервіси й не запускаються як "
        "локальні контейнери у docker-compose.yml.",
    )
    set_text(
        find_paragraph(doc, "Інтелектуальний модуль (ML / LangChain)"),
        "Сервіс ml-service приймає запит із URL зображення, виконує "
        "preprocessing та inference моделі, після чого повертає прогноз породи "
        "і confidence до серверного маршруту Next.js.",
    )
    set_text(
        find_paragraph(doc, "Для оркестрації контейнерів використовується Docker Compose"),
        "Для локального та серверного запуску використовується Docker Compose, "
        "який описує саме два реалізовані контейнери web і ml-service та "
        "передає Next.js адресу ML_SERVICE_URL. Nginx, MinIO і MLflow не "
        "належать до поточної конфігурації системи.",
    )


def update_implementation_text(doc: Document) -> None:
    insert_after(
        find_paragraph(doc, "Окремий API-маршрут /api/upload"),
        "Для визначення породи додано серверний маршрут /api/predict. "
        "Адміністративна форма передає йому URL уже завантаженого зображення "
        "та тип тварини, а маршрут викликає FastAPI-сервіс за адресою з "
        "ML_SERVICE_URL. Це усуває залежність клієнтського браузера від "
        "локальної адреси ML-сервісу й відповідає схемам інтеграції.",
    )
    set_text(
        find_paragraph(doc, "Для документування та тестування серверної частини використовується Swagger"),
        "Для документування і ручної перевірки ML-інтерфейсу використовується "
        "автоматично сформована FastAPI документація OpenAPI (Swagger UI). "
        "Маршрути основного Next.js-застосунку реалізовані як API Routes і "
        "Server Actions, а Swagger-інтерфейс на рисунку 3.3 стосується "
        "окремого FastAPI ML-сервісу.",
    )
    set_text(
        find_paragraph(doc, "Таким чином, використання Swagger"),
        "Таким чином, Swagger UI забезпечує перевірку контракту ML-сервісу та "
        "полегшує діагностику його відповіді, тоді як наскрізний сценарій "
        "застосунку перевіряється через Next.js API /api/predict.",
    )
    set_text(
        find_paragraph(doc, "Інтелектуальний модуль у системі AdoptPet"),
        "Модуль машинного навчання у системі AdoptPet реалізований як окремий "
        "FastAPI-сервіс для класифікації породи тварини за фотографією. "
        "AI-помічник на базі LangChain/OpenAI є іншим компонентом і відповідає "
        "за текстові рекомендації та генерацію опису картки.",
    )
    set_text(
        find_paragraph(doc, "Взаємодія між основною серверною частиною застосунку"),
        "Взаємодія з ML-сервісом виконується через серверний API-маршрут "
        "Next.js /api/predict. Клієнтська форма не звертається безпосередньо "
        "до FastAPI: вона надсилає imageUrl та тип тварини маршруту Next.js, "
        "а той проксіює JSON-запит до ML_SERVICE_URL/predict і повертає результат.",
    )
    set_text(
        find_paragraph(doc, "Процес взаємодії починається з отримання даних від користувача"),
        "Процес починається з вибору фотографії у формі створення поста. "
        "Компонент ImageUploadInput надсилає файл до /api/upload, де файл "
        "валідується та завантажується у Cloudinary. Після отримання imageUrl "
        "адміністратор запускає визначення породи, передаючи URL та тип тварини "
        "до /api/predict.",
    )
    set_text(
        find_paragraph(doc, "Крім того, використання API забезпечує гнучкість інтеграції"),
        "FastAPI завантажує зображення за URL, виконує preprocessing та inference "
        "відповідної моделі для кота або собаки й повертає bestPrediction та "
        "topPredictions. Під час підтвердження форми Server Action addPost "
        "зберігає публічний Post із результатами ML і створює синхронізовану "
        "внутрішню картку Animal.",
    )


def update_ml_and_deployment_text(doc: Document) -> None:
    insert_after(
        find_paragraph(doc, "Рис. 4.13."),
        "На рисунку 4.13 суцільними блоками показано компоненти, що входять "
        "до поточної реалізації: підготовку наборів даних, навчання та "
        "оцінювання моделей, збереження checkpoints і class maps, FastAPI "
        "Prediction API, preprocessing для inference та повернення top "
        "predictions. Блоки Model Registry, автоматичного моніторингу і "
        "сповіщень позначені як перспективні та не заявляються як реалізовані "
        "функції дипломного застосунку.",
    )
    set_text(
        find_paragraph(doc, "Розгортання платформи AdoptPet передбачає запуск"),
        "Розгортання платформи AdoptPet передбачає запуск fullstack-застосунку "
        "Next.js та окремого FastAPI ML-сервісу, а також підключення до "
        "зовнішніх хмарних сервісів MongoDB Atlas, Cloudinary, OpenAI API й "
        "провайдера автентифікації. Така структура відповідає фактичним "
        "компонентам системи й дозволяє ізолювати обчислення ML від веб-рівня.",
    )
    set_text(
        find_paragraph(doc, "Після запуску користувач взаємодіє з платформою через браузер"),
        "Після запуску користувач взаємодіє з платформою через браузер, а "
        "сторінки, API Routes і Server Actions обробляються Next.js. "
        "Зображення передаються через /api/upload до Cloudinary. Під час "
        "визначення породи форма звертається до /api/predict, який надсилає "
        "HTTP-запит у контейнер ml-service та повертає JSON-прогноз до інтерфейсу.",
    )


def main() -> None:
    missing = [str(path) for path in [SOURCE, *IMAGES.values()] if not path.exists()]
    if missing:
        raise FileNotFoundError("\n".join(missing))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)

    replace_diagrams(doc)
    update_architecture_text(doc)
    update_implementation_text(doc)
    update_ml_and_deployment_text(doc)

    doc.save(OUTPUT)
    print(OUTPUT.resolve())
    print(f"paragraphs={len(doc.paragraphs)}, inline_shapes={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
